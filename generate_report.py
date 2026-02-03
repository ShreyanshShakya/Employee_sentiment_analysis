# %% 
# Script to generate Final Report DOCX
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Title
title = doc.add_heading('Employee Sentiment Analysis', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Final Project Report')
doc.add_paragraph('Date: February 2026')
doc.add_paragraph()

# Executive Summary
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    'This report presents the findings from a comprehensive sentiment analysis of 2,191 employee '
    'email communications from 10 employees spanning January 2010 to December 2011. Using TextBlob '
    'for sentiment analysis, each message was classified as Positive, Negative, or Neutral based on '
    'polarity scores. The analysis reveals a balanced sentiment distribution (44.5% positive, 46.6% neutral, '
    '8.9% negative), with 6 employees identified as flight risks based on their communication patterns.'
)

# Methodology
doc.add_heading('1. Approach and Methodology', level=1)

doc.add_heading('1.1 Sentiment Labeling (TextBlob)', level=2)
doc.add_paragraph(
    'Library Used: TextBlob (rule-based sentiment analysis)\n'
    'Classification Thresholds:\n'
    '• Polarity > 0.1 → Positive (+1)\n'
    '• Polarity < -0.1 → Negative (-1)\n'
    '• Otherwise → Neutral (0)'
)

doc.add_heading('1.2 Flight Risk Criteria', level=2)
doc.add_paragraph(
    'An employee is flagged as a flight risk if they sent 4 or more negative emails within '
    'any rolling 30-day window, irrespective of calendar month boundaries.'
)

doc.add_heading('1.3 Predictive Modeling', level=2)
doc.add_paragraph(
    'Algorithm: Linear Regression (scikit-learn)\n'
    'Features: Message frequency, average message length, word count, previous month score\n'
    'Train/Test Split: 80/20'
)

# Key Findings from EDA
doc.add_heading('2. Key Findings from EDA', level=1)

doc.add_heading('2.1 Dataset Overview', level=2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
cells = table.rows[0].cells
cells[0].text = 'Metric'
cells[1].text = 'Value'
data = [
    ('Total Messages', '2,191'),
    ('Unique Employees', '10'),
    ('Date Range', 'Jan 2010 - Dec 2011'),
    ('Most Active Employee', 'Lydia Delgado (284 messages)')
]
for i, (metric, value) in enumerate(data, 1):
    table.rows[i].cells[0].text = metric
    table.rows[i].cells[1].text = value

doc.add_paragraph()
doc.add_heading('2.2 Sentiment Distribution', level=2)
table2 = doc.add_table(rows=4, cols=3)
table2.style = 'Table Grid'
table2.rows[0].cells[0].text = 'Sentiment'
table2.rows[0].cells[1].text = 'Count'
table2.rows[0].cells[2].text = 'Percentage'
sentiment_data = [
    ('Positive', '975', '44.5%'),
    ('Neutral', '1,020', '46.6%'),
    ('Negative', '196', '8.9%')
]
for i, (sent, count, pct) in enumerate(sentiment_data, 1):
    table2.rows[i].cells[0].text = sent
    table2.rows[i].cells[1].text = count
    table2.rows[i].cells[2].text = pct

# Employee Scoring and Ranking
doc.add_heading('3. Employee Scoring and Ranking', level=1)

doc.add_heading('3.1 Scoring Process', level=2)
doc.add_paragraph(
    'Monthly sentiment scores were calculated by summing individual message scores (+1, -1, 0). '
    'Average monthly score was +3.27, indicating overall positive sentiment in the workplace.'
)

doc.add_heading('3.2 Top 3 Positive Employees', level=2)
table3 = doc.add_table(rows=4, cols=3)
table3.style = 'Table Grid'
table3.rows[0].cells[0].text = 'Rank'
table3.rows[0].cells[1].text = 'Employee'
table3.rows[0].cells[2].text = 'Total Score'
pos_data = [
    ('1', 'rhonda.denton@enron.com', '50'),
    ('2', 'kayne.coulter@enron.com', '61'),
    ('3', 'bobette.riner@ipgdirect.com', '72')
]
for i, (rank, emp, score) in enumerate(pos_data, 1):
    table3.rows[i].cells[0].text = rank
    table3.rows[i].cells[1].text = emp
    table3.rows[i].cells[2].text = score

doc.add_paragraph()
doc.add_heading('3.3 Top 3 Negative Employees', level=2)
table4 = doc.add_table(rows=4, cols=3)
table4.style = 'Table Grid'
table4.rows[0].cells[0].text = 'Rank'
table4.rows[0].cells[1].text = 'Employee'
table4.rows[0].cells[2].text = 'Total Score'
neg_data = [
    ('1', 'bobette.riner@ipgdirect.com', '72'),
    ('2', 'john.arnold@enron.com', '76'),
    ('3', 'lydia.delgado@enron.com', '77')
]
for i, (rank, emp, score) in enumerate(neg_data, 1):
    table4.rows[i].cells[0].text = rank
    table4.rows[i].cells[1].text = emp
    table4.rows[i].cells[2].text = score

# Flight Risk
doc.add_heading('4. Flight Risk Identification', level=1)
doc.add_paragraph(
    'Employees were flagged as flight risks if they sent 4 or more negative emails within '
    'any 30-day rolling window. 6 out of 10 employees (60%) were identified as flight risks:'
)
doc.add_paragraph(
    '1. bobette.riner@ipgdirect.com\n'
    '2. johnny.palmer@enron.com\n'
    '3. lydia.delgado@enron.com\n'
    '4. patti.thompson@enron.com\n'
    '5. rhonda.denton@enron.com\n'
    '6. sally.beck@enron.com'
)

# Predictive Model
doc.add_heading('5. Predictive Model Overview', level=1)

doc.add_heading('5.1 Model Performance', level=2)
table5 = doc.add_table(rows=5, cols=2)
table5.style = 'Table Grid'
table5.rows[0].cells[0].text = 'Metric'
table5.rows[0].cells[1].text = 'Value'
model_data = [
    ('R² Score', '0.41'),
    ('MAE', '1.56'),
    ('RMSE', '2.17'),
    ('Training Samples', '184')
]
for i, (metric, value) in enumerate(model_data, 1):
    table5.rows[i].cells[0].text = metric
    table5.rows[i].cells[1].text = value

doc.add_paragraph()
doc.add_heading('5.2 Feature Importance', level=2)
doc.add_paragraph(
    '• Message Frequency: +0.338 (strongest positive predictor - more active = more positive)\n'
    '• Word Count: +0.031 (positive correlation)\n'
    '• Previous Month Score: -0.046 (weak negative correlation)\n'
    '• Message Length: -0.004 (negligible impact)'
)

doc.add_heading('5.3 Model Interpretation', level=2)
doc.add_paragraph(
    'The linear regression model explains 41% of variance in sentiment scores (R²=0.41). '
    'Message frequency is the strongest predictor - employees who send more messages tend to '
    'have more positive sentiment. This suggests engagement correlates with positive workplace attitude.'
)

# Conclusions
doc.add_heading('6. Conclusions and Recommendations', level=1)
doc.add_paragraph(
    '1. BALANCED SENTIMENT: Using TextBlob, 44.5% positive and 46.6% neutral messages indicate '
    'generally professional and constructive workplace communication.\n\n'
    '2. FLIGHT RISK DETECTION: 60% of employees were flagged, suggesting the 4-negative-emails '
    'threshold may be too sensitive for this dataset. Consider adjusting thresholds based on context.\n\n'
    '3. ENGAGEMENT MATTERS: Message frequency strongly predicts positive sentiment - encouraging '
    'active participation may improve overall workplace morale.\n\n'
    '4. MODEL SUCCESS: R²=0.41 is a reasonable result for sentiment prediction, showing that '
    'simple features can capture meaningful patterns in communication behavior.'
)

# Visualizations reference
doc.add_heading('7. Visualizations', level=1)
doc.add_paragraph(
    'The following visualizations are available in the visualizations/ folder:\n'
    '• sentiment_distribution.png - Sentiment breakdown pie/bar chart\n'
    '• sentiment_trends_monthly.png - Monthly trends stacked bar chart\n'
    '• sentiment_ratio_trends.png - Positive/negative percentage trends\n'
    '• top_active_employees.png - Most active employees\n'
    '• message_length_analysis.png - Message length by sentiment\n'
    '• monthly_score_distribution.png - Score distribution histogram\n'
    '• score_heatmap.png - Employee-month score heatmap\n'
    '• employee_rankings.png - Monthly rankings\n'
    '• flight_risk_analysis.png - Flight risk distribution\n'
    '• model_performance.png - Model evaluation plots'
)

# Save
doc.save('Final_Report.docx')
print('✅ Final_Report.docx created successfully with TextBlob results!')
